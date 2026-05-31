"""问题台账与报告导出（§3.6）。后端用 python-docx 生成 docx 报告。"""
from __future__ import annotations

import io
from typing import List

from sqlalchemy.orm import Session

from app.models.entities import CheckTask, Document, IssueRecord


def build_report_docx(db: Session, task: CheckTask) -> bytes:
    from docx import Document as Docx
    from docx.shared import Pt

    document: Document = task.document
    issues: List[IssueRecord] = list(task.issues)

    doc = Docx()
    doc.add_heading("文档合规检查报告", level=0)

    meta = doc.add_paragraph()
    meta.add_run("检查文件：").bold = True
    meta.add_run(document.file_name)
    doc.add_paragraph(f"文档分类：{document.category or '未分类'}")
    doc.add_paragraph(f"检查模板：{task.template_key}")
    doc.add_paragraph(f"检查结论：{task.summary}")

    doc.add_heading("问题台账", level=1)
    if not issues:
        doc.add_paragraph("本次检查未发现疑点。")
    else:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        headers = ["序号", "疑点描述", "资料位置", "法规依据", "风险等级", "整改建议"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
        for idx, it in enumerate(issues, 1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = it.description
            row[2].text = it.location or ""
            row[3].text = it.legal_basis or "—"
            row[4].text = it.risk_level or ""
            row[5].text = it.suggestion or ""

    doc.add_paragraph()
    note = doc.add_paragraph(
        "说明：本报告由合规检查智能体自动生成，作为辅助审查依据，"
        "最终结论须经人工复核确认。"
    )
    note.runs[0].font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

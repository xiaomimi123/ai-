"""核查报告生成（v3 §3.6）。

按 v3 文档规定的 5 章节结构生成 Word 报告：
封面
第一部分：总体核查结论
第二部分：佐证材料总体合规性核查
第三部分：各指标相关性核查明细
第四部分：高频风险点汇总
第五部分：整改建议清单
"""
from __future__ import annotations

import io
import json
from collections import Counter, defaultdict
from datetime import datetime
from typing import List, Optional

from sqlalchemy.orm import Session

from app.models import AuditTask, AuditUnit, Finding, Indicator, Material, Worksheet, WorksheetRow
from app.services.scoring_service import compute_task_scoring


def build_report_docx(db: Session, task: AuditTask) -> bytes:
    """生成 Word 核查报告，返回 docx 二进制内容。"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    unit: AuditUnit = task.unit
    materials: List[Material] = list(task.materials)
    findings: List[Finding] = list(task.findings)
    scoring = compute_task_scoring(db, task)

    # 拉工作底稿（V2：审计师在底稿编辑的内容会注入报告）
    worksheet: Optional[Worksheet] = (
        db.query(Worksheet).filter(Worksheet.task_id == task.id).first()
    )
    ws_rows_by_ind: dict[int, WorksheetRow] = {}
    if worksheet:
        for row in worksheet.rows:
            if row.indicator_id:
                ws_rows_by_ind[row.indicator_id] = row

    doc = Document()

    # 中文字体设置
    def _set_cn_font(run, name="宋体", size=12, bold=False):
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)

    def _h(text, level=1, size=18, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        r = p.add_run(text)
        _set_cn_font(r, name="黑体", size=size, bold=True)
        return p

    def _p(text, size=11, indent=True, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        _set_cn_font(r, size=size)
        return p

    def _label_value(label, value):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}：")
        _set_cn_font(r1, name="黑体", size=11, bold=True)
        r2 = p.add_run(str(value))
        _set_cn_font(r2, size=11)
        return p

    # ============================================================
    # 封面
    # ============================================================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title.add_run("内部控制评价核查报告")
    _set_cn_font(rt, name="黑体", size=28, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = subtitle.add_run(f"{task.eval_year} 年度")
    _set_cn_font(rs, name="楷体", size=18)

    for _ in range(8):
        doc.add_paragraph()

    cover_meta = [
        ("被检查单位", unit.name if unit else "—"),
        ("评价年度", f"{task.eval_year} 年度"),
        ("核查任务", task.name),
        ("核查日期", _fmt_date(task.completed_at or task.created_at)),
        ("核查依据", "评价指标库、问题清单库、上位法、编报指南"),
        ("报告生成时间", _fmt_date(datetime.utcnow())),
    ]
    for label, value in cover_meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}：")
        _set_cn_font(r1, name="黑体", size=12, bold=True)
        r2 = p.add_run(str(value))
        _set_cn_font(r2, size=12)

    doc.add_page_break()

    # ============================================================
    # 第一部分 总体核查结论
    # ============================================================
    _h("第一部分 · 总体核查结论", size=18)

    conclusion = _build_conclusion(findings)
    _p(f"本次核查覆盖被检查单位提交的 {len(materials)} 份佐证材料，"
       f"涉及 {_count_indicators(materials)} 个评价指标，"
       f"AI 共检出 {len(findings)} 条疑点。")
    _label_value("总体结论", conclusion)

    # ─── 评分卡片（功能 3 新增）───
    if scoring and scoring["total_max"] > 0:
        doc.add_paragraph()
        _h("综合评分", level=2, size=14)
        score_para = doc.add_paragraph()
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = score_para.add_run(f"{scoring['total_score']}")
        _set_cn_font(r1, name="黑体", size=40, bold=True)
        r2 = score_para.add_run(f"  /  {scoring['total_max']} 分")
        _set_cn_font(r2, name="黑体", size=18)

        grade_para = doc.add_paragraph()
        grade_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = grade_para.add_run(f"得分率 {scoring['score_pct']}%　·　等级")
        _set_cn_font(r3, size=12)
        r4 = grade_para.add_run(f"「{scoring['grade']}」")
        _set_cn_font(r4, name="黑体", size=14, bold=True)
        doc.add_paragraph()
        _p(
            f"评分规则：高风险扣指标满分 50%，中风险扣 25%，低风险扣 10%；"
            f"已忽略的发现不扣分，已调整的按 50% 计。"
            f"等级阈值：优 ≥90 / 良 ≥80 / 中 ≥60 / 差 <60。",
            size=9,
        )

    # 风险分布表
    _h("问题分布统计", level=2, size=14)
    sev_counter = Counter(f.severity for f in findings)
    type_counter = Counter(f.finding_type for f in findings)

    sev_table = doc.add_table(rows=2, cols=4)
    sev_table.style = "Light Grid Accent 1"
    headers = ["风险等级", "高", "中", "低"]
    counts = ["数量", str(sev_counter.get("高", 0)),
              str(sev_counter.get("中", 0)),
              str(sev_counter.get("低", 0))]
    for i, h in enumerate(headers):
        cell = sev_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                _set_cn_font(r, name="黑体", size=11, bold=True)
    for i, c in enumerate(counts):
        cell = sev_table.rows[1].cells[i]
        cell.text = c
        for p in cell.paragraphs:
            for r in p.runs:
                _set_cn_font(r, size=11)

    doc.add_paragraph()

    # 按类型分布
    if type_counter:
        _h("按问题类型分布", level=2, size=14)
        type_table = doc.add_table(rows=1 + len(type_counter), cols=2)
        type_table.style = "Light Grid Accent 1"
        for i, hdr in enumerate(["问题类型", "数量"]):
            cell = type_table.rows[0].cells[i]
            cell.text = hdr
            for p in cell.paragraphs:
                for r in p.runs:
                    _set_cn_font(r, name="黑体", size=11, bold=True)
        for row_i, (k, v) in enumerate(sorted(type_counter.items(), key=lambda x: -x[1]), start=1):
            type_table.rows[row_i].cells[0].text = k
            type_table.rows[row_i].cells[1].text = str(v)
            for cell in type_table.rows[row_i].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        _set_cn_font(r, size=11)

    doc.add_paragraph()
    _p("（说明：本报告由 AI 核查引擎自动生成，作为辅助审查依据，"
       "最终结论须经人工复核确认。AI 核查的是文本内容合规性，"
       "原件真实性需结合现场核查或档案比对确认。）", size=10)

    # ─── 评分汇总表（按指标）───
    if scoring and scoring["indicators"]:
        doc.add_paragraph()
        _h("评分汇总表（按指标）", level=2, size=14)
        rows_data = scoring["indicators"]
        # 表头 + 数据行
        table = doc.add_table(rows=1 + len(rows_data) + 1, cols=6)
        table.style = "Light Grid Accent 1"
        headers = ["指标编号", "名称", "满分", "扣分", "得分", "问题数"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    _set_cn_font(r, name="黑体", size=10, bold=True)
        for row_i, item in enumerate(rows_data, 1):
            row = table.rows[row_i]
            row.cells[0].text = item["indicator_code"]
            row.cells[1].text = _truncate(item["name"], 30)
            row.cells[2].text = str(item["max_score"])
            row.cells[3].text = str(item["deducted"])
            row.cells[4].text = str(item["actual_score"])
            row.cells[5].text = str(item["findings_total"])
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        _set_cn_font(r, size=10)
        # 合计行
        total_row = table.rows[-1]
        total_row.cells[0].text = "合计"
        total_row.cells[2].text = str(scoring["total_max"])
        total_row.cells[3].text = str(scoring["total_deducted"])
        total_row.cells[4].text = str(scoring["total_score"])
        total_row.cells[5].text = str(sum(it["findings_total"] for it in rows_data))
        for cell in total_row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    _set_cn_font(r, name="黑体", size=10, bold=True)

    doc.add_page_break()

    # ============================================================
    # 第二部分 佐证材料总体合规性核查
    # ============================================================
    _h("第二部分 · 佐证材料总体合规性核查", size=18)
    _p("本部分针对所有佐证材料的真实性、年度一致性、正式性、要素完整性，"
       "按 4 个维度汇总核查结果。")

    overall_dims = {
        "真实性": ["真实性问题"],
        "年度一致性": ["年度一致性问题"],
        "正式性": ["正式性问题"],
        "要素完整性": ["完整性问题"],
    }

    for dim_name, types in overall_dims.items():
        dim_findings = [f for f in findings if f.finding_type in types]
        _h(f"二·{list(overall_dims.keys()).index(dim_name)+1}  {dim_name}核查结果",
           level=2, size=13)
        if not dim_findings:
            _p(f"✓ 本维度未发现问题。", size=11)
        else:
            _p(f"共发现 {len(dim_findings)} 条问题：", size=11)
            _add_finding_table(doc, dim_findings, _set_cn_font)
        doc.add_paragraph()

    doc.add_page_break()

    # ============================================================
    # 第三部分 各指标核查明细（按指标分组，每条 finding 5 段式呈现）
    # ============================================================
    _h("第三部分 · 各指标核查明细", size=18)
    _p("按评价指标分组，对每条问题明确列出：评价标准（指标要求）/ 不符合之处 / "
       "调整建议 / 法规依据 / 风险等级及整改时限。")

    # 按 indicator 分组所有 finding（不限类型，让所有指标级问题都体现）
    grouped: dict[int, list[Finding]] = defaultdict(list)
    no_ind_findings = []
    for f in findings:
        if f.indicator_id:
            grouped[f.indicator_id].append(f)
        else:
            no_ind_findings.append(f)

    if not grouped and not no_ind_findings:
        _p("✓ 各指标未发现问题。")
    else:
        for ind_id, ind_findings in grouped.items():
            ind = db.get(Indicator, ind_id)
            if ind is None:
                continue
            # 指标小标题
            # 从评分明细里找该指标的得分
            ind_score = next((s for s in scoring.get("indicators", [])
                              if s["indicator_id"] == ind.id), None)
            score_suffix = ""
            if ind_score:
                score_suffix = f"  （得分 {ind_score['actual_score']} / {ind_score['max_score']} 分）"
            _h(f"指标 {ind.indicator_code} · {ind.name}{score_suffix}", level=2, size=13)
            _label_value("分类", f"{ind.level} / {ind.category} / {ind.subcategory}")
            _label_value("满分", str(ind.max_score))

            # V2/V3：优先用工作底稿的审计师评分 / 评语
            ws_row = ws_rows_by_ind.get(ind.id)
            if ws_row:
                _label_value("核查前得分（自评）", f"{ws_row.original_score} 分")
                _label_value("核查后得分（复核）", f"{ws_row.audited_score} 分")
                # 新版优先用 adjustment_note（调整得分说明）
                note = ws_row.adjustment_note or ws_row.audit_finding_text
                if note:
                    _label_value("审计师说明", note)
            else:
                if ind_score:
                    _label_value("扣分", f"{ind_score['deducted']} 分")
                    _label_value("实际得分", f"{ind_score['actual_score']} 分")
            _label_value("问题数", f"{len(ind_findings)} 条")
            doc.add_paragraph()

            # 按风险排序
            sev_order = {"高": 0, "中": 1, "低": 2}
            ind_findings.sort(key=lambda x: sev_order.get(x.severity, 9))
            for sub_idx, f in enumerate(ind_findings, 1):
                _add_finding_detailed(doc, sub_idx, f, ind, _set_cn_font, _p, _label_value)

            doc.add_paragraph()
            # 章节分隔线
            sep = doc.add_paragraph("―" * 30)
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in sep.runs:
                _set_cn_font(r, size=10)
            doc.add_paragraph()

        # 无指标关联的 finding（共享池或异常）
        if no_ind_findings:
            _h("其它问题（未关联具体指标）", level=2, size=13)
            for sub_idx, f in enumerate(no_ind_findings, 1):
                _add_finding_detailed(doc, sub_idx, f, None, _set_cn_font, _p, _label_value)

    doc.add_page_break()

    # ============================================================
    # 第四部分 高频风险点汇总
    # ============================================================
    _h("第四部分 · 高频风险点汇总", size=18)
    _p("按风险级别 + 类型聚合，提供整改的优先级参考。")

    high_findings = [f for f in findings if f.severity == "高"]
    _h(f"四·1  高风险问题（共 {len(high_findings)} 条）", level=2, size=13)
    if not high_findings:
        _p("✓ 未发现高风险问题。")
    else:
        _add_finding_table(doc, high_findings, _set_cn_font)

    doc.add_paragraph()

    # 真实性 + 完整性 高频
    real_findings = [f for f in findings
                     if f.finding_type in ("真实性问题", "完整性问题")
                     and f.severity in ("高", "中")]
    _h(f"四·2  真实性与完整性问题（共 {len(real_findings)} 条）", level=2, size=13)
    if not real_findings:
        _p("✓ 未发现此类问题。")
    else:
        _add_finding_table(doc, real_findings, _set_cn_font)

    doc.add_page_break()

    # ============================================================
    # 第五部分 整改建议清单
    # ============================================================
    _h("第五部分 · 整改建议清单", size=18)
    _p("按风险等级排序，列出每条问题的整改建议、建议完成时限。")

    # 排序：高 > 中 > 低
    sev_order = {"高": 0, "中": 1, "低": 2}
    sorted_findings = sorted(findings, key=lambda f: (sev_order.get(f.severity, 9), f.id))
    if not sorted_findings:
        _p("✓ 本次核查未发现问题。")
    else:
        suggestion_table = doc.add_table(rows=1 + len(sorted_findings), cols=5)
        suggestion_table.style = "Light Grid Accent 1"
        for i, hdr in enumerate(["序号", "风险", "问题摘要", "整改建议", "建议时限"]):
            cell = suggestion_table.rows[0].cells[i]
            cell.text = hdr
            for p in cell.paragraphs:
                for r in p.runs:
                    _set_cn_font(r, name="黑体", size=10, bold=True)
        for idx, f in enumerate(sorted_findings, 1):
            row = suggestion_table.rows[idx]
            row.cells[0].text = str(idx)
            row.cells[1].text = f.severity
            row.cells[2].text = _truncate(f.description, 80)
            row.cells[3].text = _truncate(f.suggestion or "—", 100)
            row.cells[4].text = _suggested_deadline(f.severity)
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        _set_cn_font(r, size=10)

    doc.add_paragraph()
    _p("（建议时限说明：高风险问题 30 日内整改完成；"
       "中风险问题 90 日内整改完成；低风险问题可在下次评价周期内整改。）",
       size=10)

    # ============================================================
    # 签章页
    # ============================================================
    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = foot.add_run("核查方（盖章）：__________________________")
    _set_cn_font(r, name="黑体", size=12)
    for _ in range(3):
        doc.add_paragraph()
    foot2 = doc.add_paragraph()
    foot2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = foot2.add_run(f"报告生成日期：{_fmt_date(datetime.utcnow())}")
    _set_cn_font(r, size=12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================
# Helpers
# ============================================================
def _build_conclusion(findings: List[Finding]) -> str:
    if not findings:
        return "通过 — 核查未发现问题"
    high = sum(1 for f in findings if f.severity == "高")
    if high == 0:
        return f"有条件通过 — 共 {len(findings)} 条疑点，需关注并整改"
    return f"不通过 — 检出 {high} 条高风险问题，需重点整改"


def _count_indicators(materials: List[Material]) -> int:
    return len({m.indicator_id for m in materials if m.indicator_id})


def _add_finding_table(doc, findings: List[Finding], set_font):
    """通用：把一组 finding 渲染成表格（简版，用于第二部分等汇总场景）。"""
    if not findings:
        return
    table = doc.add_table(rows=1 + len(findings), cols=5)
    table.style = "Light Grid Accent 1"
    for i, hdr in enumerate(["序号", "风险", "问题描述", "材料位置", "建议"]):
        cell = table.rows[0].cells[i]
        cell.text = hdr
        for p in cell.paragraphs:
            for r in p.runs:
                set_font(r, name="黑体", size=10, bold=True)
    for idx, f in enumerate(findings, 1):
        row = table.rows[idx]
        row.cells[0].text = str(idx)
        row.cells[1].text = f.severity
        row.cells[2].text = _truncate(f.description, 100)
        row.cells[3].text = _truncate(f.evidence_location or "—", 60)
        row.cells[4].text = _truncate(f.suggestion or "—", 80)
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_font(r, size=10)


def _add_finding_detailed(doc, idx: int, f: Finding, indicator,
                          set_cn_font, p_helper, label_value):
    """详细版：每条 finding 用 5 段式呈现 — 评价标准 / 不符合之处 / 调整建议 / 法规依据 / 风险与时限。

    用于第三部分（各指标核查明细）。
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # 标题行：[X] 不符合：{指标编号} {指标名称}
    title = doc.add_paragraph()
    title.paragraph_format.space_before = doc.styles["Normal"].paragraph_format.space_before
    r0 = title.add_run(f"[{idx}] ")
    set_cn_font(r0, name="黑体", size=11, bold=True)
    chip = title.add_run(f"风险 {f.severity}")
    set_cn_font(chip, name="黑体", size=10, bold=True)
    sep = title.add_run("  ·  ")
    set_cn_font(sep, size=10)
    if indicator is not None:
        t2 = title.add_run(f"{indicator.indicator_code} {indicator.name}")
        set_cn_font(t2, name="黑体", size=11, bold=True)
    else:
        t2 = title.add_run(f.finding_type)
        set_cn_font(t2, name="黑体", size=11, bold=True)

    # ▸ 评价标准（来自 indicator.description / deduct_rules）
    if indicator:
        std = indicator.description or indicator.name
        if indicator.deduct_rules:
            std += f"（扣分细则：{_truncate(indicator.deduct_rules, 200)}）"
        _bullet(doc, "评价标准", std, set_cn_font)

    # ▸ 不符合之处（finding.description + evidence_location）
    not_ok = f.description
    if f.evidence_location:
        not_ok += f"（位置：{f.evidence_location}）"
    _bullet(doc, "不符合之处", not_ok, set_cn_font, highlight=True)

    # ▸ 调整建议
    if f.suggestion:
        _bullet(doc, "调整建议", f.suggestion, set_cn_font)

    # ▸ 法规依据（来自 RAG 召回 / indicator 扣分细则）
    if f.legal_basis:
        _bullet(doc, "法规依据", _truncate(f.legal_basis, 600), set_cn_font)
    elif indicator and indicator.deduct_rules:
        _bullet(doc, "法规依据", indicator.deduct_rules, set_cn_font)

    # ▸ 风险等级 + 整改时限
    _bullet(doc, "风险等级",
            f"{f.severity}　|　建议整改时限：{_suggested_deadline(f.severity)}",
            set_cn_font)

    doc.add_paragraph()  # 间隔


def _bullet(doc, label, value, set_cn_font, highlight=False):
    """5 段式中的一段：▸ 标签：值"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = doc.styles["Normal"].paragraph_format.left_indent
    from docx.shared import Cm
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = None
    arrow = p.add_run("▸ ")
    set_cn_font(arrow, size=10, bold=True)
    lbl = p.add_run(f"{label}：")
    set_cn_font(lbl, name="黑体", size=10, bold=True)
    val = p.add_run(str(value))
    set_cn_font(val, size=10, bold=highlight)


def _truncate(text: str, n: int) -> str:
    if not text:
        return "—"
    text = str(text)
    return text if len(text) <= n else text[:n] + "…"


def _fmt_date(dt) -> str:
    if dt is None:
        return "—"
    if isinstance(dt, str):
        return dt
    try:
        return dt.strftime("%Y 年 %m 月 %d 日")
    except Exception:
        return str(dt)


def _suggested_deadline(severity: str) -> str:
    return {"高": "30 日内", "中": "90 日内", "低": "下次评价周期"}.get(severity, "—")

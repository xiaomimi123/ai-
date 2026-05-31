"""Word(.docx) 解析器（python-docx）。"""
from __future__ import annotations

from pathlib import Path

from app.parsers.base import ParsedDocument, PageBlock
from app.parsers.txt_parser import _SECTION_RE


def parse_docx(path: str) -> ParsedDocument:
    from docx import Document  # 延迟导入，未安装时仅影响该格式

    doc = Document(path)
    blocks: list[PageBlock] = []
    texts: list[str] = []
    current_section = "正文"

    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        if para.style and para.style.name and para.style.name.lower().startswith("heading"):
            current_section = t[:40]
        elif _SECTION_RE.match(t):
            current_section = t[:40]
        blocks.append(PageBlock(page=1, section=current_section, content=t))
        texts.append(t)

    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tables.append(rows)
        # 表格文本也并入全文，供规则与检索
        for row in rows:
            texts.append("\t".join(row))

    return ParsedDocument(
        text="\n".join(texts),
        page_blocks=blocks,
        tables=tables,
        metadata={"file_name": Path(path).name, "parser": "docx"},
    )
